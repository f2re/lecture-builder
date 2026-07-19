from __future__ import annotations

import re
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Cm
from lxml import etree

from .models import ValidationResult

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
NS = {"w": W_NS, "m": M_NS}
EQ_NUMBER_RE = re.compile(r"^\((\d+)\.(\d+)\)$")


def _close(a: int | None, b: int, tolerance: int = 1500) -> bool:
    return a is not None and abs(a - b) <= tolerance


def _attribute(element, local_name: str) -> str | None:
    if element is None:
        return None
    return element.get(f"{{{W_NS}}}{local_name}")


def _table_layout_is_valid(table) -> tuple[bool, str | None]:
    cells = table.xpath("./w:tr[1]/w:tc", namespaces=NS)
    if len(cells) != 3:
        return False, "equation table must have exactly three cells"
    formula_cells = [cell for cell in cells if cell.xpath(".//m:oMath | .//m:oMathPara", namespaces=NS)]
    number_cells = [
        cell
        for cell in cells
        if any(EQ_NUMBER_RE.match("".join(cell.itertext()).strip()) for _ in [0])
    ]
    if len(formula_cells) != 1 or len(number_cells) != 1:
        return False, "equation table must contain one formula cell and one number cell"
    formula_alignment = formula_cells[0].xpath(".//w:pPr/w:jc[1]", namespaces=NS)
    number_alignment = number_cells[0].xpath(".//w:pPr/w:jc[1]", namespaces=NS)
    if not formula_alignment or _attribute(formula_alignment[0], "val") != "center":
        return False, "formula cell is not centered"
    if not number_alignment or _attribute(number_alignment[0], "val") != "right":
        return False, "equation number is not right-aligned"

    borders = table.xpath("./w:tblPr/w:tblBorders/*", namespaces=NS)
    if borders and any(_attribute(border, "val") not in {"nil", "none"} for border in borders):
        return False, "equation table has visible borders"
    return True, None


def validate_docx(path: str | Path, *, expect_formulas: bool | None = None) -> ValidationResult:
    source = Path(path)
    result = ValidationResult(name="docx")
    if not source.is_file() or source.stat().st_size == 0:
        result.add("docx.missing", "DOCX отсутствует или пуст", path=source)
        return result

    try:
        document = Document(str(source))
    except Exception as exc:
        result.add("docx.open", f"Не удалось открыть DOCX: {exc}", path=source)
        return result

    section = document.sections[0]
    page_expected = {"width": Cm(21.0).emu, "height": Cm(29.7).emu}
    page_actual = {"width": section.page_width, "height": section.page_height}
    for name, expected_value in page_expected.items():
        actual_value = int(page_actual[name]) if page_actual[name] is not None else None
        if not _close(actual_value, expected_value):
            result.add(
                "docx.page_size",
                f"Размер страницы {name} не соответствует A4",
                path=source,
                details={"expected_emu": expected_value, "actual_emu": actual_value},
            )

    expected = {
        "left": Cm(3.0).emu,
        "right": Cm(1.5).emu,
        "top": Cm(2.0).emu,
        "bottom": Cm(2.0).emu,
    }
    actual = {
        "left": section.left_margin,
        "right": section.right_margin,
        "top": section.top_margin,
        "bottom": section.bottom_margin,
    }
    for name, expected_value in expected.items():
        actual_value = int(actual[name]) if actual[name] is not None else None
        if not _close(actual_value, expected_value):
            result.add(
                "docx.margin",
                f"Поле {name} не соответствует профилю ГОСТ",
                path=source,
                details={"expected_emu": expected_value, "actual_emu": actual_value},
            )

    normal = document.styles["Normal"]
    if normal.font.name != "Times New Roman":
        result.add(
            "docx.font",
            f"Стиль Normal использует шрифт {normal.font.name!r} вместо Times New Roman",
            path=source,
        )
    if normal.font.size is None or abs(normal.font.size.pt - 14) > 0.2:
        result.add(
            "docx.font_size",
            f"Размер Normal равен {normal.font.size.pt if normal.font.size else None} pt вместо 14 pt",
            path=source,
        )

    try:
        with zipfile.ZipFile(source) as archive:
            document_xml = archive.read("word/document.xml")
            footer_xml = b"".join(
                archive.read(name)
                for name in archive.namelist()
                if name.startswith("word/footer") and name.endswith(".xml")
            )
    except (zipfile.BadZipFile, KeyError, OSError) as exc:
        result.add("docx.zip", f"Некорректная структура DOCX: {exc}", path=source)
        return result

    root = etree.fromstring(document_xml)
    math_count = len(root.xpath(".//m:oMath | .//m:oMathPara", namespaces=NS))
    marker_count = len(root.xpath(".//w:t[contains(., '[[EQNO:')]", namespaces=NS))
    equation_tables = 0
    valid_equation_tables = 0
    equation_numbers: list[str] = []
    for table in root.xpath(".//w:tbl", namespaces=NS):
        texts = ["".join(node.itertext()).strip() for node in table.xpath(".//w:tc", namespaces=NS)]
        numbers = [text for text in texts if EQ_NUMBER_RE.match(text)]
        has_math = bool(table.xpath(".//m:oMath | .//m:oMathPara", namespaces=NS))
        if numbers and has_math:
            equation_tables += 1
            equation_numbers.extend(numbers)
            valid, reason = _table_layout_is_valid(table)
            if valid:
                valid_equation_tables += 1
            else:
                result.add(
                    "docx.equation_alignment",
                    f"Некорректное оформление формулы: {reason}",
                    path=source,
                )

    if marker_count:
        result.add("docx.equation_marker", "В DOCX остались технические маркеры EQNO", path=source)
    if expect_formulas is True and math_count == 0:
        result.add("docx.omml", "Ожидались нативные формулы OMML, но они не найдены", path=source)
    if expect_formulas is True and equation_tables == 0:
        result.add(
            "docx.equation_layout",
            "Не найдены безрамочные таблицы с формулой и номером справа",
            path=source,
        )
    if equation_tables and valid_equation_tables != equation_tables:
        result.add(
            "docx.equation_layout_count",
            "Не все пронумерованные формулы имеют корректное выравнивание и границы",
            path=source,
            details={"total": equation_tables, "valid": valid_equation_tables},
        )

    parsed = [EQ_NUMBER_RE.match(number) for number in equation_numbers]
    prefixes = [int(match.group(1)) for match in parsed if match]
    ordinals = [int(match.group(2)) for match in parsed if match]
    if prefixes and len(set(prefixes)) != 1:
        result.add("docx.equation_prefix", "В DOCX смешаны номера разных лекций", path=source)
    if ordinals and ordinals != list(range(1, len(ordinals) + 1)):
        result.add(
            "docx.equation_sequence",
            f"Номера формул в DOCX непоследовательны: {ordinals}",
            path=source,
        )

    footer_text = footer_xml.decode("utf-8", errors="ignore")
    if "PAGE" not in footer_text:
        result.add("docx.page_number", "В нижнем колонтитуле отсутствует поле PAGE", path=source)

    result.metrics = {
        "size_bytes": source.stat().st_size,
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "omml_nodes": math_count,
        "numbered_equation_tables": equation_tables,
        "valid_numbered_equation_tables": valid_equation_tables,
    }
    return result
