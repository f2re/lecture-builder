from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from .constants import (
    BODY_SIZE,
    CAPTION_SIZE,
    CODE_SIZE,
    FIRST_LINE,
    FONT,
    FONT_CODE,
    HEADING_1_SIZE,
    HEADING_SIZE,
    LINE_SPACING,
)


def set_font(style, name: str, size: int, *, bold: bool = False, italic: bool = False) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)


def set_paragraph_style(
    style,
    *,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line=FIRST_LINE,
    before: int = 0,
    after: int = 0,
    line_spacing=LINE_SPACING,
) -> None:
    paragraph = style.paragraph_format
    paragraph.alignment = alignment
    paragraph.first_line_indent = first_line
    paragraph.space_before = Pt(before)
    paragraph.space_after = Pt(after)
    paragraph.line_spacing = line_spacing
    paragraph.line_spacing_rule = WD_LINE_SPACING.EXACTLY


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, end))
    run.font.name = FONT
    run.font.size = Pt(CAPTION_SIZE)


def configure_section(section) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)


def build_reference(path: Path) -> None:
    document = Document()
    configure_section(document.sections[0])

    normal = document.styles["Normal"]
    set_font(normal, FONT, BODY_SIZE)
    set_paragraph_style(normal)

    heading1 = document.styles["Heading 1"]
    set_font(heading1, FONT, HEADING_1_SIZE, bold=True)
    set_paragraph_style(
        heading1,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line=Cm(0),
        before=12,
        after=8,
    )
    heading1.paragraph_format.keep_with_next = True

    for name, level in (("Heading 2", 2), ("Heading 3", 3), ("Heading 4", 4)):
        style = document.styles[name]
        set_font(style, FONT, HEADING_SIZE, bold=(level <= 3))
        set_paragraph_style(
            style,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line=Cm(0),
            before=10 if level == 2 else 8,
            after=4,
        )
        style.paragraph_format.keep_with_next = True

    for name in ("Body Text", "First Paragraph", "List Paragraph"):
        try:
            style = document.styles[name]
        except KeyError:
            style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        set_font(style, FONT, BODY_SIZE)
        set_paragraph_style(style, first_line=FIRST_LINE if name != "List Paragraph" else Cm(0))

    try:
        caption = document.styles["Caption"]
    except KeyError:
        caption = document.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    set_font(caption, FONT, CAPTION_SIZE)
    set_paragraph_style(
        caption,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line=Cm(0),
        before=3,
        after=6,
        line_spacing=Pt(14),
    )

    try:
        block = document.styles["Block Text"]
    except KeyError:
        block = document.styles.add_style("Block Text", WD_STYLE_TYPE.PARAGRAPH)
    set_font(block, FONT, CAPTION_SIZE)
    set_paragraph_style(block, first_line=Cm(0), line_spacing=Pt(15))
    block.paragraph_format.left_indent = Cm(1.25)

    for name in ("Verbatim", "Source Code"):
        try:
            style = document.styles[name]
        except KeyError:
            continue
        set_font(style, FONT_CODE, CODE_SIZE)
        set_paragraph_style(style, first_line=Cm(0), line_spacing=Pt(14))

    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.clear()
    add_page_field(paragraph)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)
