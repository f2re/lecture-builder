#!/usr/bin/env python3
"""
md2docx_gost.py — Markdown → DOCX (ГОСТ 7.32, нативные формулы Word)

Зависимости:
    pip install pypandoc python-docx lxml pandoc
"""

import argparse, re, sys, tempfile
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ─── Параметры ГОСТ ───────────────────────────────────────────────────────────
FONT        = "Times New Roman"
FONT_CODE   = "Courier New"
SZ_BODY     = 14;  SZ_H1 = 16;  SZ_H2 = 14;  SZ_CAP = 12;  SZ_CODE = 12
LINE_SP     = Pt(18)            # точный межстрочный интервал
INDENT      = Cm(1.25)          # красная строка
REF_NAME    = "gost_reference.docx"


# ════════════════════════════════════════════════════════════════════════════════
# ШАГ 1 — Создание reference.docx со стилями ГОСТ
# ════════════════════════════════════════════════════════════════════════════════
def build_reference(path: Path):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width   = Cm(21);   sec.page_height  = Cm(29.7)
    sec.left_margin  = Cm(3.0);  sec.right_margin = Cm(1.5)   # ГОСТ 30/15 мм
    sec.top_margin   = Cm(2.0);  sec.bottom_margin= Cm(2.0)

    def font(style, name, size, bold=False, italic=False):
        f = style.font
        f.name, f.size, f.bold, f.italic = name, Pt(size), bold, italic
        rpr = style.element.get_or_add_rPr()
        rf  = rpr.find(qn('w:rFonts'))
        if rf is None:
            rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
        for a in ('w:ascii','w:hAnsi','w:cs','w:eastAsia'):
            rf.set(qn(a), name)

    def para(style, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first=INDENT,
             sb=0, sa=0, ls=LINE_SP):
        pf = style.paragraph_format
        pf.alignment = align;  pf.first_line_indent = first
        pf.space_before = Pt(sb);  pf.space_after = Pt(sa)
        pf.line_spacing = ls;  pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    # Normal (Обычный)
    font(doc.styles['Normal'], FONT, SZ_BODY);  para(doc.styles['Normal'])

    # Heading 1 — по центру, жирный 16пт
    font(doc.styles['Heading 1'], FONT, SZ_H1, bold=True)
    para(doc.styles['Heading 1'], WD_ALIGN_PARAGRAPH.CENTER, Cm(0), sb=12, sa=6)
    doc.styles['Heading 1'].paragraph_format.keep_with_next = True

    # Heading 2 — слева, жирный 14пт
    font(doc.styles['Heading 2'], FONT, SZ_H2, bold=True)
    para(doc.styles['Heading 2'], WD_ALIGN_PARAGRAPH.LEFT, Cm(0), sb=10, sa=4)
    doc.styles['Heading 2'].paragraph_format.keep_with_next = True

    # Heading 3 — с красной строкой, обычный
    font(doc.styles['Heading 3'], FONT, SZ_H2)
    para(doc.styles['Heading 3'], WD_ALIGN_PARAGRAPH.LEFT, INDENT, sb=8, sa=4)

    # Body Text и First Paragraph
    for sname in ('Body Text', 'First Paragraph'):
        try:   st = doc.styles[sname]
        except KeyError:
               st = doc.styles.add_style(sname, WD_STYLE_TYPE.PARAGRAPH)
        font(st, FONT, SZ_BODY);  para(st)

    # Block Text (цитаты)
    try:   bt = doc.styles['Block Text']
    except KeyError:
           bt = doc.styles.add_style('Block Text', WD_STYLE_TYPE.PARAGRAPH)
    font(bt, FONT, SZ_CAP);  para(bt, first=Cm(0))
    bt.paragraph_format.left_indent = INDENT

    # Caption (подписи к рисункам, таблицам)
    try:   cap = doc.styles['Caption']
    except KeyError:
           cap = doc.styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
    font(cap, FONT, SZ_CAP)
    para(cap, WD_ALIGN_PARAGRAPH.CENTER, Cm(0), sb=3, sa=6)

    # Verbatim (блоки кода)
    for sname in ('Verbatim', 'Source Code', 'Code'):
        try:   cs = doc.styles[sname]
        except KeyError:
            try: cs = doc.styles.add_style(sname, WD_STYLE_TYPE.PARAGRAPH)
            except: continue
        font(cs, FONT_CODE, SZ_CODE);  para(cs, first=Cm(0))
        cs.paragraph_format.left_indent = Cm(1.0)

    # Нижний колонтитул — номер страницы
    footer = sec.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    for kind, txt in [('begin',None),(None,'PAGE'),('end',None)]:
        if kind:
            el = OxmlElement('w:fldChar'); el.set(qn('w:fldCharType'), kind)
        else:
            el = OxmlElement('w:instrText'); el.text = txt
        run._r.append(el)
    run.font.name = FONT;  run.font.size = Pt(SZ_CAP)

    doc.save(str(path))
    print(f"[✓] reference.docx → {path}")


# ════════════════════════════════════════════════════════════════════════════════
# ШАГ 2 — Предобработка Markdown
# ════════════════════════════════════════════════════════════════════════════════
def preprocess(src: Path) -> str:
    text = src.read_text(encoding='utf-8')
    text = re.sub(r'^---\n.*?---\n', '', text, flags=re.DOTALL)  # YAML
    text = re.sub(r'<!--.*?-->', '', text, flags=re.DOTALL)       # HTML-комментарии
    
    # \( ... \) → $ ... $ (inline math)
    text = re.sub(r'\\\(\s*(.*?)\s*\\\)', r'$\1$', text, flags=re.DOTALL)
    # \[ ... \] → $$ ... $$ (display math)
    text = re.sub(r'\\\[\s*(.*?)\s*\\\]', r'$$\1$$', text, flags=re.DOTALL)
    
    # Удаляем \tag{...} из формул, так как pandoc (mathml) их не понимает
    text = re.sub(r'\\tag\{.*?\}', '', text)
    
    # Кавычки-елочки: заменяем прямые кавычки на типографские
    text = re.sub(r'(^|[\s\(\[«])"', r'\1«', text)
    text = re.sub(r'"', r'»', text)
    
    text = re.sub(r'\$\$\s*\$\$', '', text)  # пустые блоки
    return text


# ════════════════════════════════════════════════════════════════════════════════
# ШАГ 3 — Конвертация через pandoc
# ════════════════════════════════════════════════════════════════════════════════
def convert(md_text: str, ref: Path, out: Path, toc: bool):
    # Ищем pandoc в нескольких местах
    import shutil, os
    
    # Список возможных путей
    search_paths = [
        shutil.which('pandoc'),
        os.path.expanduser('~/bin/pandoc'),
        '/home/user/bin/pandoc',
        '/usr/bin/pandoc',
        '/usr/local/bin/pandoc'
    ]
    
    pandoc_bin = None
    for p in search_paths:
        if p and os.path.exists(p) and os.access(p, os.X_OK):
            pandoc_bin = p
            break
            
    if not pandoc_bin:
        print("❌ ОШИБКА: pandoc не найден. Установите его: sudo apt install pandoc", file=sys.stderr)
        sys.exit(1)

    with tempfile.NamedTemporaryFile(mode='w', suffix='.md',
                                     delete=False, encoding='utf-8') as tf:
        tf.write(md_text);  tmp = tf.name

    cmd = [
        pandoc_bin, tmp,
        '-f', 'markdown+tex_math_dollars+tex_math_double_backslash',
        '-t', 'docx',
        '--mathml',                      # ← LaTeX → нативные уравнения Word
        f'--reference-doc={ref}',
        '--wrap=none',
        '--standalone',
        '-o', str(out),
    ]
    if toc:
        cmd += ['--toc', '--toc-depth=3']

    import subprocess
    r = subprocess.run(cmd, capture_output=True, text=True)
    Path(tmp).unlink(missing_ok=True)

    if r.returncode != 0:
        print(f"[pandoc stderr]\n{r.stderr}", file=sys.stderr)
        sys.exit(r.returncode)
    print(f"[✓] pandoc → {out}")


# ════════════════════════════════════════════════════════════════════════════════
# ШАГ 4 — Постобработка python-docx
# ════════════════════════════════════════════════════════════════════════════════
def postprocess(path: Path):
    doc = Document(str(path))

    def fix(run, name=FONT, size=SZ_BODY):
        run.font.name = name
        if not run.font.size or run.font.size < Pt(8):
            run.font.size = Pt(size)
        rpr = run._r.get_or_add_rPr()
        rf  = rpr.find(qn('w:rFonts'))
        if rf is None:
            rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
        for a in ('w:ascii','w:hAnsi','w:cs','w:eastAsia'):
            rf.set(qn(a), name)

    for p in doc.paragraphs:
        sn = p.style.name
        # Блоки с формулами — по центру
        if not p.text.strip() and p._p.find(qn('m:oMath')) is not None:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(6)
            continue
        sz = SZ_H1 if 'Heading 1' in sn else SZ_H2 if 'Heading' in sn else SZ_BODY
        fn = FONT_CODE if sn in ('Verbatim','Source Code','Code') else FONT
        for run in p.runs:
            fix(run, fn, SZ_CODE if fn == FONT_CODE else sz)

    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        fix(run, FONT, SZ_CAP)
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after  = Pt(2)

    doc.save(str(path))
    print(f"[✓] postprocess → {path}")


# ════════════════════════════════════════════════════════════════════════════════
# MAIN
# ════════════════════════════════════════════════════════════════════════════════
def main():
    ap = argparse.ArgumentParser(
        description='Конвертер Markdown → DOCX (ГОСТ, нативные формулы Word)')
    ap.add_argument('input')
    ap.add_argument('-o', '--output', default=None)
    ap.add_argument('--toc',         action='store_true', default=True)
    ap.add_argument('--no-toc',      action='store_true')
    ap.add_argument('--rebuild-ref', action='store_true')
    ap.add_argument('--ref-docx',    default=None)
    args = ap.parse_args()

    src = Path(args.input).resolve()
    if not src.exists():
        sys.exit(f"Файл не найден: {src}")

    out = Path(args.output).resolve() if args.output else src.with_suffix('.docx')
    ref = Path(args.ref_docx).resolve() if args.ref_docx \
          else Path(__file__).parent / REF_NAME

    if args.rebuild_ref or not ref.exists():
        build_reference(ref)
    else:
        print(f"[i] reference.docx: {ref}")

    md  = preprocess(src)
    toc = args.toc and not args.no_toc
    convert(md, ref, out, toc)
    postprocess(out)

    print(f"\n✅ Готово: {out}  ({out.stat().st_size // 1024} КБ)")


if __name__ == '__main__':
    main()
