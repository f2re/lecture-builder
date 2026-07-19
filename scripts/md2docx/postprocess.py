from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from .constants import (
    BODY_SIZE,
    CAPTION_SIZE,
    CODE_SIZE,
    FONT,
    FONT_CODE,
    HEADING_1_SIZE,
    HEADING_SIZE,
)
from .reference import add_page_field, configure_section

MARKER_RE = re.compile(r"^\s*\[\[EQNO:(?P<number>\d+\.\d+)\]\]\s*$")


def paragraph_has_math(paragraph_element) -> bool:
    return bool(paragraph_element.xpath(".//m:oMath | .//m:oMathPara"))


def remove_table_borders(table) -> None:
    properties = table._tbl.tblPr
    borders = properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "nil")


def set_cell_width(cell, width: Cm) -> None:
    cell.width = width
    properties = cell._tc.get_or_add_tcPr()
    tcw = properties.first_child_found_in("w:tcW")
    if tcw is None:
        tcw = OxmlElement("w:tcW")
        properties.append(tcw)
    tcw.set(qn("w:w"), str(int(width.twips)))
    tcw.set(qn("w:type"), "dxa")


def configure_run(run, *, name: str = FONT, size: int = BODY_SIZE) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)


def find_formula_before(marker_paragraph):
    node = marker_paragraph._p.getprevious()
    while node is not None:
        if node.tag == qn("w:p"):
            if paragraph_has_math(node):
                return node
            text = "".join(node.itertext()).strip()
            if text:
                return None
        node = node.getprevious()
    return None


def insert_equation_rows(document: Document) -> int:
    converted = 0
    for marker in list(document.paragraphs):
        match = MARKER_RE.match(marker.text)
        if not match:
            continue
        number = match.group("number")
        formula_element = find_formula_before(marker)
        if formula_element is None:
            raise RuntimeError(f"Не найдена OMML-формула перед маркером {number}")

        parent = formula_element.getparent()
        insertion_index = parent.index(formula_element)

        table = document.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        remove_table_borders(table)
        cells = table.rows[0].cells
        widths = (Cm(2.0), Cm(12.5), Cm(2.0))
        for cell, width in zip(cells, widths, strict=True):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)

        table_element = table._tbl
        table_element.getparent().remove(table_element)
        parent.insert(insertion_index, table_element)

        middle = cells[1]
        placeholder = middle.paragraphs[0]._p
        middle._tc.remove(placeholder)
        formula_element.getparent().remove(formula_element)
        middle._tc.append(formula_element)
        formula_paragraph = middle.paragraphs[0]
        formula_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        formula_paragraph.paragraph_format.first_line_indent = Cm(0)
        formula_paragraph.paragraph_format.space_before = Pt(4)
        formula_paragraph.paragraph_format.space_after = Pt(4)

        right = cells[2].paragraphs[0]
        right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right.paragraph_format.first_line_indent = Cm(0)
        right.paragraph_format.space_before = Pt(0)
        right.paragraph_format.space_after = Pt(0)
        run = right.add_run(f"({number})")
        configure_run(run)

        marker._p.getparent().remove(marker._p)
        converted += 1
    return converted


def postprocess(output: Path, expected_equations: list[str]) -> None:
    document = Document(output)
    for index, section in enumerate(document.sections):
        configure_section(section)
        footer = section.footer
        if index == 0:
            footer.is_linked_to_previous = False
            paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            paragraph.clear()
            add_page_field(paragraph)
        else:
            footer.is_linked_to_previous = True

    converted = insert_equation_rows(document)
    if converted != len(expected_equations):
        raise RuntimeError(
            f"Число оформленных формул ({converted}) не совпадает с числом тегов ({len(expected_equations)})"
        )

    for paragraph in document.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        paragraph.paragraph_format.widow_control = True
        if style_name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.first_line_indent = Cm(0)
        font_name = FONT_CODE if style_name in {"Verbatim", "Source Code", "Code"} else FONT
        font_size = CODE_SIZE if font_name == FONT_CODE else (
            HEADING_1_SIZE
            if style_name == "Heading 1"
            else HEADING_SIZE
            if style_name.startswith("Heading")
            else BODY_SIZE
        )
        for run in paragraph.runs:
            configure_run(run, name=font_name, size=font_size)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        configure_run(
                            run,
                            size=CAPTION_SIZE if not paragraph_has_math(paragraph._p) else BODY_SIZE,
                        )

    document.save(output)
